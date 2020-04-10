from django.shortcuts import render
from django.views.generic.base import TemplateView
from .forms import ResumeForm
from django.urls import reverse_lazy
from django.views.generic import DetailView, ListView
from django.http import HttpResponseRedirect
from .models import Resume
import os
from docx import Document


from pyresparser import ResumeParser

class ResumeUploadView(TemplateView):
    resume_form = ResumeForm
    template_name = "index.html"

    def post(self, request):

        post_data = request.POST or None
        file_data = request.FILES or None

        resume_form = ResumeForm(post_data, file_data)

        if resume_form.is_valid():
            resume_obj = resume_form.save()

            data = ResumeParser('media/'+str(resume_obj.document)).get_extracted_data()

            resume_obj.name = data.get('name') or 'Not Found'
            resume_obj.email = data.get('email') or 'Not Found'
            resume_obj.phone_number = data.get('mobile_number') or 'Not Found'

            file_name, file_extension = os.path.splitext('media/'+str(resume_obj.document))

            from custom_resume_parser.resume_parser import CustomResumeParser
            data11 = CustomResumeParser.get_extracted_data('media/'+str(resume_obj.document))
            print(data11)

            if file_extension == '.docx':
                font_style, font_size, total_tables, total_images, total_char, total_lines, linkedin_url, phone, email = extract_from_docx('media/'+str(resume_obj.document))

                if phone:
                    resume_obj.phone_number = phone
                if email:
                    resume_obj.email  = email

                resume_obj.font_name = font_style
                resume_obj.font_size = font_size[:2]
                resume_obj.total_tables = total_tables
                resume_obj.total_images = total_images
                resume_obj.total_char = total_char
                resume_obj.total_lines = total_lines
                resume_obj.linkedin_url = linkedin_url or 'Not Found'

            elif file_extension == '.pdf':



                resume_obj.font_name = 'Arial'
                resume_obj.font_size = '24'
                resume_obj.total_tables = 3
                resume_obj.total_images = 5
                resume_obj.total_char = 30
                resume_obj.total_lines = 120
                resume_obj.linkedin_url = 'https://www.linkedin.com/'
                pass

            
            resume_obj.save()
            
            return HttpResponseRedirect(reverse_lazy('resumedetail', kwargs={'pk': resume_obj.id}))

        context = self.get_context_data(
                                        resume_form=resume_form,
                                    )

        return self.render_to_response(context)     

    def get(self, request, *args, **kwargs):
        return self.post(request, *args, **kwargs)



class ResumeDetailView(DetailView):
    model = Resume
    template_name = "resumedetail.html"
    context_object_name = 'resume'

class ResumeListView(ListView):
    model = Resume
    template_name = "resumelist.html"
    context_object_name = 'resumes'


def extract_from_docx(path):
    f = open(path, 'rb')
    document = Document(f)
    f.close()

    style_list = []
    char_count = 0
    total_line_count = 0
    phone = ''
    email = ''

    import re
    reg_link = re.compile('^(http(s?)://)?([\w\d\-\_]*\.)?linkedin.com/[^.\s]*$')
    reg_phone = re.compile('^(\+\d{1,2}\s)?\(?\d{3}\)?[\s.-]\d{3}[\s.-]\d{4}$')
    reg_email = re.compile('^\w+([\.-]?\w+)*@\w+([\.-]?\w+)*(\.\w{2,3})+$')

    for p in document.paragraphs:
        name = p.style.font.name
        size = p.style.font.size
        char_count += len(p.text)
        style_list.append([name, size])
        x = p.text
        total_line_count += len(x.split('\n'))
        if reg_phone.search(x):
            phone = x
        if reg_email.search(x):
            email = x
   
    import itertools
    style_list.sort()
    style_list1 = list(k for k,_ in itertools.groupby(style_list))
    font_style = ''
    font_size = ''
    for style in style_list1:
        font_style += (str(style[0]) + ', ' )
        font_size += (str(style[1]) + ', ' )
    
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    rels = document.part.rels
    def iter_hyperlink_rels(rels):
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                jj = reg_link.search(str(rels[rel]._target))
                if jj:
                    return str(rels[rel]._target)
    linkedin = iter_hyperlink_rels(rels)

    return font_style, font_size, len(document.tables), len(document.inline_shapes), char_count, total_line_count, linkedin, phone, email


