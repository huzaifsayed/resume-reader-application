from docx import Document
import itertools

def get_font_table_shape_from_docx(path):
    f = open(path, 'rb')
    document = Document(f)
    f.close()

    style_list = []

    for p in document.paragraphs:
        name = p.style.font.name
        size = p.style.font.size
        style_list.append([name, size])
   
    style_list.sort()
    style_list1 = list(k for k,_ in itertools.groupby(style_list))
    font_style = ''
    font_size = ''
    for style in style_list1:
        font_style += (str(style[0]) + ', ' )
        font_size += (str(style[1]) + ', ' )

    return {
                'font_style': font_style, 
                'font_size': font_size, 
                'doc_table': len(document.tables),
                'doc_shapes': len(document.inline_shapes),
            }

# try:
#     from xml.etree.cElementTree import XML
# except ImportError:
#     from xml.etree.ElementTree import XML
# import zipfile


# """
# Module that extract text from MS XML Word document (.docx).
# (Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
# """

# WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
# PARA = WORD_NAMESPACE + 'ascii'
# TEXT = WORD_NAMESPACE + 't'


# def get_docx_text(path):
#     """
#     Take the path of a docx file as argument, return the text in unicode.
#     """
#     document = zipfile.ZipFile(path)
#     xml_content = document.read('word/document.xml')
#     document.close()
#     tree = XML(xml_content)

#     paragraphs = []
#     for paragraph in tree.getiterator(PARA):
#         texts = [node.text
#                  for node in paragraph.getiterator(TEXT)
#                  if node.text]
#         if texts:
#             paragraphs.append(''.join(texts))

#     return '\n\n'.join(paragraphs)